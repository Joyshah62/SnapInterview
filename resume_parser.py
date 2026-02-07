#!/usr/bin/env python3
"""
Resume parser: uses parser.GeneralResumeParser (pypdf) for PDFs and python-docx for DOCX.
No dependency on unstructured.
"""
from pathlib import Path


def _parse_pdf(file_path: str) -> str:
    """Parse PDF using GeneralResumeParser (parser.py)."""
    from parser import GeneralResumeParser
    parser = GeneralResumeParser(file_path)
    return parser.parse(output_file=None, include_raw=False)


def _parse_docx(file_path: str) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for .docx resumes. Install with: pip install python-docx"
        )
    doc = Document(file_path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                parts.append(" | ".join(row_text))
    return "\n".join(parts) if parts else ""


def parse_document(file_path: str) -> str:
    """
    Parse resume document (PDF or DOCX).
    Returns extracted text. Uses parser.GeneralResumeParser for PDF, python-docx for DOCX.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _parse_pdf(file_path)
    if suffix in (".docx", ".doc"):
        return _parse_docx(file_path)
    # Fallback: try PDF parser for unknown extension (e.g. upload without extension)
    if suffix in ("", ".bin"):
        try:
            return _parse_pdf(file_path)
        except Exception:
            pass
    raise ValueError(
        f"Unsupported file type: {suffix}. Use .pdf or .docx"
    )


def parse_and_save(file_path: str, output_path: str = "resume_text.md") -> str:
    """
    Parse resume and save to output_path. Same API as before for server/interview flow.
    """
    text = parse_document(file_path)
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(text, encoding="utf-8")
    return text


def main():
    import sys
    input_file = sys.argv[1] if len(sys.argv) > 1 else "document.pdf"
    output_file = sys.argv[2] if len(sys.argv) > 2 else "resume_text.md"
    try:
        doc_text = parse_and_save(input_file, output_file)
        print(f"Saved to {output_file}")
    except Exception as e:
        print(f"\n❌ Error: {e}")
        print("\nPDF support: pypdf (parser.py uses it)")
        print("DOCX support: pip install python-docx")


if __name__ == "__main__":
    main()
